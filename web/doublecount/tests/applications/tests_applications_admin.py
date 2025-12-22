from datetime import date
from io import BytesIO

from django.contrib.auth import get_user_model
from django.test import TestCase
from django.urls import reverse
from docx import Document

from certificates.models import DoubleCountingRegistration
from core.models import Entity, ExternalAdminRights, MatierePremiere, Pays, UserRights
from core.tests_utils import setup_current_user
from doublecount.factories import (
    DoubleCountingApplicationFactory,
    DoubleCountingProductionFactory,
    DoubleCountingSourcingFactory,
)
from doublecount.models import DoubleCountingApplication, DoubleCountingProduction
from doublecount.views.applications.mixins.approve_application import DoubleCountingApplicationApproveError
from doublecount.views.applications.mixins.utils import (
    DoubleCountingApplicationExportError,
    application_to_json,
    check_has_dechets_industriels,
)
from transactions.models import ProductionSite


class Endpoint:
    change_user_role = reverse("api-entity-users-change-role")


User = get_user_model()


class AdminDoubleCountApplicationsTest(TestCase):
    fixtures = [
        "json/biofuels.json",
        "json/feedstock.json",
        "json/countries.json",
        "json/depots.json",
        "json/entities.json",
        "json/entities_sites.json",
    ]

    def setUp(self):
        self.admin = Entity.objects.filter(entity_type=Entity.ADMIN)[0]
        self.ext_admin = Entity.objects.create(name="ExternalAdminTest", entity_type=Entity.EXTERNAL_ADMIN)
        ExternalAdminRights.objects.create(entity=self.ext_admin, right=ExternalAdminRights.DOUBLE_COUNTING)

        self.user = setup_current_user(
            self, "tester@carbure.local", "Tester", "gogogo", [(self.admin, "RW"), (self.ext_admin, "RW")], True
        )

        self.producer, _ = Entity.objects.update_or_create(name="Le Super Producteur 1", entity_type="Producteur")
        UserRights.objects.update_or_create(user=self.user, entity=self.producer, defaults={"role": UserRights.ADMIN})

        self.production_site = ProductionSite.objects.first()
        self.production_site.created_by = self.producer
        self.production_site.address = "1 rue de la Paix"
        france, _ = Pays.objects.update_or_create(code_pays="FR", name="France")
        self.production_site.country = france
        self.production_site.city = "Paris"
        self.production_site.postal_code = "75000"
        self.production_site.dc_reference = ""
        self.production_site.save()
        self.requested_start_year = 2023

    def create_application(self):
        app = DoubleCountingApplicationFactory.create(
            producer=self.production_site.producer,
            production_site=self.production_site,
            period_start=date(self.requested_start_year, 1, 1),
            period_end=date(self.requested_start_year + 1, 12, 31),
            status=DoubleCountingApplication.PENDING,
        )

        sourcing1 = DoubleCountingSourcingFactory.create(dca=app, year=self.requested_start_year)
        sourcing2 = DoubleCountingSourcingFactory.create(dca=app, year=self.requested_start_year)
        production1 = DoubleCountingProductionFactory.create(
            dca=app, feedstock=sourcing2.feedstock, year=self.requested_start_year + 1, approved_quota=-1
        )
        production2 = DoubleCountingProductionFactory.create(
            dca=app, feedstock=sourcing2.feedstock, year=self.requested_start_year + 1, approved_quota=-1
        )

        return app, sourcing1, production1, sourcing1, production2

    def test_list_applications(self):
        self.create_application()

        response = self.client.get(
            reverse("double-counting-applications-list-admin"),
            {"entity_id": self.admin.id, "year": self.requested_start_year},
        )

        data = response.json()
        pending = data["pending"]
        application = pending[0]

        assert application["producer"]["id"] == self.production_site.producer.id

    def test_list_applications_for_external_admin(self):
        self.create_application()

        response = self.client.get(
            reverse("double-counting-applications-list-admin"),
            {"entity_id": self.ext_admin.id, "year": self.requested_start_year},
        )

        application = response.json()["pending"][0]
        assert application["producer"]["id"] == self.production_site.producer.id

    def test_application_details(self):
        app, sourcing, production, _, _ = self.create_application()

        response = self.client.get(
            reverse("double-counting-applications-detail", kwargs={"id": app.id}),
            {"entity_id": self.admin.id, "dca_id": app.id},
        )

        application = response.json()

        assert application["sourcing"][0]["id"] == sourcing.id
        assert application["production"][0]["id"] == production.id

        production_site = application["production_site"]
        assert production_site["id"] == self.production_site.id
        assert production_site["inputs"] == []
        assert production_site["certificates"] == []

    def test_approve_application(self):
        application, sourcing1, production1, sourcing2, production2 = self.create_application()

        # test approve without quotas
        params = {"dca_id": application.id, "entity_id": self.admin.id}
        response = self.client.post(reverse("double-counting-applications-approve"), params)
        assert response.status_code == 400
        assert response.json()["message"] == DoubleCountingApplicationApproveError.QUOTAS_NOT_APPROVED
        assert application.status == DoubleCountingApplication.PENDING

        # update quotas
        updated_quotas = [[production1.id, 20500], [production2.id, 10000]]
        response = self.client.post(
            reverse(
                "double-counting-applications-update-approved-quotas",
                kwargs={"id": application.id},
            )
            + f"?entity_id={self.admin.id}",
            data={"approved_quotas": updated_quotas},
            content_type="application/json",
        )
        assert response.status_code == 200

        productions = DoubleCountingProduction.objects.filter(dca_id=application.id)
        assert productions[0].approved_quota == 20500
        assert productions[1].approved_quota == 10000

        application = DoubleCountingApplication.objects.get(
            producer=self.production_site.producer, period_start__year=self.requested_start_year
        )

        # test approve with quotas
        response = self.client.post(reverse("double-counting-applications-approve"), params)
        assert response.status_code == 200
        application = DoubleCountingApplication.objects.get(
            producer=self.production_site.producer, period_start__year=self.requested_start_year
        )
        assert application.status == DoubleCountingApplication.ACCEPTED

        # test agreement generation
        agreement = DoubleCountingRegistration.objects.get(certificate_id=application.certificate_id)
        assert agreement.valid_from == date(2023, 1, 1)
        assert agreement.valid_until == date(2024, 12, 31)
        assert agreement.production_site == application.production_site
        assert agreement.certificate_id == application.certificate_id
        assert agreement.application.id == application.id

    def test_export_application(self):
        application, sourcing1, production1, sourcing2, production2 = self.create_application()

        assert application.status != DoubleCountingApplication.ACCEPTED

        # Malformed params
        response = self.client.get(reverse("double-counting-applications-generate-decision"), {"entity_id": self.admin.id})
        assert response.status_code == 400
        assert response.json()["message"] == DoubleCountingApplicationExportError.MALFORMED_PARAMS

        # Application not found
        response = self.client.get(
            reverse("double-counting-applications-generate-decision"),
            {"dca_id": application.id + 200, "entity_id": self.admin.id},
        )
        assert response.status_code == 400
        assert response.json()["message"] == DoubleCountingApplicationExportError.APPLICATION_NOT_FOUND

        application.status = DoubleCountingApplication.ACCEPTED
        application.save()

        assert application.status == DoubleCountingApplication.ACCEPTED

        # Di without di in application
        response = self.client.get(
            reverse("double-counting-applications-generate-decision"),
            {"dca_id": application.id, "entity_id": self.admin.id, "di": "Graisses brunes, huiles acides"},
        )
        assert response.status_code == 400
        assert response.json()["message"] == DoubleCountingApplicationExportError.MALFORMED_PARAMS

        # Export without di
        assert not check_has_dechets_industriels(application)
        response = self.client.get(
            reverse("double-counting-applications-generate-decision"),
            {"dca_id": application.id, "entity_id": self.admin.id},
        )
        assert response.status_code == 200
        assert response["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # ... check doc content
        application_data = application_to_json(application)

        file_stream = BytesIO(response.content)
        doc = Document(file_stream)

        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        full_text = "\n".join(full_text)

        # ... footer
        section = doc.sections[0]
        footer = section.footer

        footer_text = []
        for paragraph in footer.paragraphs:
            footer_text.append(paragraph.text)

        footer_text = "\n".join(footer_text)
        assert application_data["certificate_id"] == footer_text

        # ... content
        del application_data["has_dechets_industriels"]
        del application_data["id"]
        del application_data["dechets_industriels"]

        for _, value in application_data.items():
            assert str(value) in full_text

        assert "Article 5" not in full_text
        assert full_text.count("Article 3") == 1

        # Export with di
        feedstock = MatierePremiere.objects.get(code="DECHETS_INDUSTRIELS")
        assert production1.feedstock.code != "DECHETS_INDUSTRIELS"
        production1.feedstock = feedstock
        production1.save()

        production1.refresh_from_db()

        assert check_has_dechets_industriels(application)

        response = self.client.get(
            reverse("double-counting-applications-generate-decision"),
            {"dca_id": application.id, "entity_id": self.admin.id, "di": "Graisses brunes, huiles acides"},
        )
        assert response.status_code == 200
        assert response["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        application_data = application_to_json(application)

        file_stream = BytesIO(response.content)
        doc = Document(file_stream)

        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        full_text = "\n".join(full_text)

        # ... content
        del application_data["has_dechets_industriels"]
        del application_data["id"]

        for _, value in application_data.items():
            assert str(value) in full_text

        assert "Article 5" in full_text
        assert full_text.count("Article 3") == 1
