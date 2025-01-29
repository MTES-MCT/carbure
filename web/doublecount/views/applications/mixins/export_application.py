import os
from io import BytesIO

from django.conf import settings
from django.db.models import Sum
from django.db.models.query_utils import Q
from django.http import HttpResponse
from docx import Document
from drf_spectacular.utils import (
    OpenApiExample,
    OpenApiParameter,
    OpenApiTypes,
    extend_schema,
)
from rest_framework import status
from rest_framework.decorators import action
from rest_framework.response import Response

from doublecount.models import DoubleCountingApplication

from .utils import (
    DoubleCountingApplicationExportError,
    WordKey,
    application_to_json,
    check_has_dechets_industriels,
    delete_paragraph,
    format_to_word,
    name_to_dc_decision_name,
    replace_and_bold,
    set_bold_text,
    set_cell_border,
    set_font,
)


class ExportApplicationActionMixin:
    @extend_schema(
        parameters=[
            OpenApiParameter(
                "entity_id",
                OpenApiTypes.INT,
                OpenApiParameter.QUERY,
                description="Entity ID",
                required=True,
            ),
            OpenApiParameter(
                "dca_id",
                OpenApiTypes.INT,
                OpenApiParameter.QUERY,
                description="Doublecount application ID",
                required=True,
            ),
            OpenApiParameter(
                "di",
                OpenApiTypes.STR,
                OpenApiParameter.QUERY,
                description="Dechet industriel",
                required=False,
            ),
        ],
        examples=[
            OpenApiExample(
                "Example of export response.",
                value="file.docx",
                request_only=False,
                response_only=True,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        ],
        responses={
            (
                200,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ): OpenApiTypes.STR,
        },
    )
    @action(methods=["get"], detail=False, url_path="export-application")
    def export_application(self, request, *args, **kwargs):
        application_id = request.query_params.get("dca_id", False)
        dechets_industriels = request.query_params.get("di", "")

        if not application_id:
            return Response(
                {"message": DoubleCountingApplicationExportError.MALFORMED_PARAMS},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            application = DoubleCountingApplication.objects.get(id=application_id)
        except Exception:
            return Response(
                {"message": DoubleCountingApplicationExportError.APPLICATION_NOT_FOUND},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # if application.status != DoubleCountingApplication.ACCEPTED:
        #     return Response(
        #         {"message": DoubleCountingApplicationExportError.APPLICATION_NOT_ACCEPTED},
        #         status=status.HTTP_400_BAD_REQUEST,
        #     )

        has_dechets_industriels = check_has_dechets_industriels(application)
        dechets_industriels = dechets_industriels.strip()
        if "," in dechets_industriels:
            part_1, part_2 = dechets_industriels.rsplit(",", 1)
            dechets_industriels = f"{part_1} et{part_2}"

        if has_dechets_industriels and not dechets_industriels:
            return Response(
                {"message": DoubleCountingApplicationExportError.DECHETS_INDUSTRIELS_NOT_FOUND},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if dechets_industriels and not has_dechets_industriels:
            return Response(
                {"message": DoubleCountingApplicationExportError.MALFORMED_PARAMS},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not dechets_industriels:
            dechets_industriels = "-"

        year_n = application.period_start.year
        year_n_1 = year_n + 1

        dcp_data = (
            application.production.filter(year__in=[year_n, year_n_1])
            .values("biofuel__name", "feedstock__name")
            .annotate(
                approved_quota_year_n=Sum("approved_quota", filter=Q(year=year_n)),
                approved_quota_year_n_1=Sum("approved_quota", filter=Q(year=year_n_1)),
            )
        )

        data = application_to_json(application, dechets_industriels)

        to_word = format_to_word(data)

        table_data = [
            {
                **item,
                "approved_quota_year_n_1": (
                    "Aucun" if item["approved_quota_year_n_1"] == 0 else item["approved_quota_year_n_1"]
                ),
            }
            for item in dcp_data
        ]

        path = os.path.join(
            settings.BASE_DIR,
            "templates",
            "word",
            "double_count_approval_template.docx",
        )
        doc = Document(path)

        # Tables
        table = doc.tables[1]
        # ... header
        cell = table.rows[0].cells
        cell[2].text = to_word.get(WordKey.YEAR_N)
        cell[3].text = to_word.get(WordKey.YEAR_N_1)

        for cell in table.rows[0].cells:
            set_font(cell.paragraphs[0])

        # ... row
        for item in table_data:
            cell = table.add_row().cells
            cell[0].text = name_to_dc_decision_name.get(item.get("biofuel__name"), (item.get("biofuel__name"), None))[0]
            cell[1].text = item.get("feedstock__name")
            cell[2].text = str(item.get("approved_quota_year_n"))
            cell[3].text = str(item.get("approved_quota_year_n_1"))

            for c in cell:
                set_font(c.paragraphs[0])
                set_cell_border(c)

        # Content
        for index, paragraph in enumerate(doc.paragraphs):
            for k, v in to_word.items():
                if k in paragraph.text:
                    target = "«DECHETS_INDUSTRIELS»"
                    if k != target:
                        paragraph.text = paragraph.text.replace(k, v)
                    elif target in paragraph.text:
                        replace_and_bold(paragraph, target, v)

            set_font(paragraph, bold=index in [5, 6, 7, 8])

        # Footer
        section = doc.sections[0]
        footer = section.footer
        for paragraph in footer.paragraphs:
            for k, v in to_word.items():
                if k in paragraph.text:
                    paragraph.text = paragraph.text.replace(k, v)

            set_font(paragraph)

        # Articles
        found_first_article_3 = False
        found_first_article_5 = False
        ARTICLE_DECHETS_INDUSTRIELS = 3
        ARTICLE_DECHETS_INDUSTRIELS_2 = 4
        current_article_number = 1
        i = 0
        while i < len(doc.paragraphs):
            paragraph = doc.paragraphs[i]

            if paragraph.text.startswith("Article"):
                article_number = int("".join(filter(str.isdigit, paragraph.text)))

                if article_number == ARTICLE_DECHETS_INDUSTRIELS:
                    if not has_dechets_industriels and not found_first_article_3:
                        delete_paragraph(doc.paragraphs[i])
                        delete_paragraph(doc.paragraphs[i])
                        found_first_article_3 = True
                        i -= 1
                    else:
                        set_bold_text(paragraph, f"Article {current_article_number}")
                        current_article_number = 4
                elif article_number == ARTICLE_DECHETS_INDUSTRIELS_2:
                    if not has_dechets_industriels and not found_first_article_5:
                        delete_paragraph(doc.paragraphs[i])
                        delete_paragraph(doc.paragraphs[i])
                        found_first_article_5 = True
                        current_article_number = 3
                        i -= 1
                    else:
                        set_bold_text(paragraph, f"Article {current_article_number}")
                        current_article_number += 1
                else:
                    set_bold_text(paragraph, f"Article {current_article_number}")
                    if current_article_number == 1:
                        superscript_run = paragraph.add_run("er")
                        superscript_run.font.superscript = True
                    current_article_number += 1

            set_font(paragraph)

            i += 1

        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f"attachment; filename={application.certificate_id}.docx"

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        response.write(doc_io.read())

        return response
