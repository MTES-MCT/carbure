from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

name_to_dc_decision_name = {
    "Éthanol pour ED9": ("Éthanol", "Éthanols", None),
    "EEHA": ("Ester Éthylique d'Huiles Animales", "Esters Éthyliques d'Huiles Animales", "Esters Éthyliques"),
    "EEHU": ("Ester Éthylique d'Huiles Usagées", "Esters Éthyliques d'Huiles Usagées", "Esters Éthyliques"),
    "EEHV": ("Ester Éthylique d'Huiles Végétales", "Esters Éthyliques d'Huiles Végétales", "Esters Éthyliques"),
    "EMAG de POME": ("Ester Méthylique d'Acides Gras", "Esters Méthyliques d'Acides Gras", "Esters Méthyliques"),
    "EMAG": ("Ester Méthylique d'Acides Gras", "Esters Méthyliques d'Acides Gras", "Esters Méthyliques"),
    "EMHA": ("Ester Méthylique d'Huiles Animales", "Esters Méthyliques d'Huiles Animales", "Esters Méthyliques"),
    "EMHU": ("Ester Méthylique d'Huiles Usagées", "Esters Méthyliques d'Huiles Usagées", "Esters Méthyliques"),
    "EMHV": ("Ester Méthylique d'Huiles Végétales", "Esters Méthyliques d'Huiles Végétales", "Esters Méthyliques"),
    "ETBE": ("Éthyl Tert-Butyl Éther", "Éthyl Tert-Butyl Éthers", "Esters Méthyliques"),
    "Éthanol": ("Éthanol", "Éthanols", "Éthanols"),
    "Huiles co-traitées - Kérosène": (
        "Huiles co-traitées carburéacteurs",
        "Huiles co-traitées carburéacteurs",
        "Huiles co-traitées",
    ),
    "Huile cotraitée - Carburéacteur": (
        "Huiles co-traitées carburéacteurs",
        "Huiles co-traitées carburéacteurs",
        "Huiles co-traitées",
    ),  # 2nd version
    "Huiles co-traitées - Essence": ("Huiles co-traitées essences", "Huiles co-traitées essences", "Huiles co-traitées"),
    "Huile cotraitée - Essence": (
        "Huiles co-traitées essences",
        "Huiles co-traitées essences",
        "Huiles co-traitées",
    ),  # 2nd version
    "Huiles co-traitées - Gazole": ("Huiles co-traitées gazoles", "Huiles co-traitées gazoles", "Huiles co-traitées"),
    "Huile cotraitée - Gazole": (
        "Huiles co-traitées gazoles",
        "Huiles co-traitées gazoles",
        "Huiles co-traitées",
    ),  # 2nd version
    "Autres Huiles Hydrotraitées - Kérosène": (
        "Huiles hydrotraitées carburéacteurs",
        "Huiles hydrotraitées carburéacteurs",
        "Huiles hydrotraitées",
    ),
    "Autres Huiles Hydrotraitées - Essence": (
        "Huiles hydrotraitées essences",
        "Huiles hydrotraitées essences",
        "Huiles hydrotraitées",
    ),
    "Autres Huiles Hydrotraitées - Gazole": (
        "Huiles hydrotraitées gazoles",
        "Huiles hydrotraitées gazoles",
        "Huiles hydrotraitées",
    ),
    "Huiles Végétales Hydrotraitées - Kérosène": (
        "Huiles Végétales hydrotraitées carburéacteurs",
        "Huiles Végétales hydrotraitées carburéacteurs",
        "Huiles hydrotraitées",
    ),
    "Huiles Végétales Hydrotraitées - Essence": (
        "Huiles Végétales hydrotraitées essences",
        "Huiles Végétales hydrotraitées essences",
        "Huiles hydrotraitées",
    ),
    "Huiles Végétales Hydrotraitées - Gazole": (
        "Huiles Végétales hydrotraitées gazoles",
        "Huiles Végétales hydrotraitées gazoles",
        "Huiles hydrotraitées",
    ),
    "Méthanol": ("Méthanol", "Méthanols", None),
    "MTBE": ("Méthyl Tert-Butyl Éther", "Méthyl Tert-Butyl Éthers", None),
    "TAEE": ("Tert-Amyl Éthyl Éther", "Tert-Amyl Éthyl Éthers", None),
    "TAME": ("Tert-Amyl Méthyl Éther", "Tert-Amyl Méthyl Éthers", None),
}


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._element = paragraph._p = None


def replace_and_bold(paragraph, old_text, new_text, font_name="Times New Roman", font_size=12):
    for run in paragraph.runs:
        if old_text in run.text:
            parts = run.text.split(old_text)
            run.text = parts[0]

            new_run = paragraph.add_run(new_text)
            new_run.font.bold = True
            new_run.font.name = font_name
            new_run.font.size = Pt(font_size)
            new_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

            if len(parts) > 1:
                remaining_run = paragraph.add_run(parts[1])
                remaining_run.font.name = font_name
                remaining_run.font.size = Pt(font_size)
                remaining_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_font(paragraph, font_name="Times New Roman", font_size=12, bold=None):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = bold


def set_bold_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True


def set_cell_border(cell, border_color="000000", border_size="4"):
    tc_pr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for border_name in ["top", "bottom", "left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), border_size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color)
        borders.append(border)
    tc_pr.append(borders)


class DoubleCountingApplicationExportError:
    MALFORMED_PARAMS = "MALFORMED_PARAMS"
    APPLICATION_NOT_FOUND = "APPLICATION_NOT_FOUND"
    APPLICATION_NOT_ACCEPTED = "APPLICATION_NOT_ACCEPTED"
    DECHETS_INDUSTRIELS_NOT_FOUND = "DECHETS_INDUSTRIELS_NOT_FOUND"


class WordKey:
    CITY = "«Ville»"
    COUNTRY = "«Pays»"
    OPERATOR_NAME = "«Nom_Opérateur»"
    ADDRESS = "«Adresse»"
    POSTAL_CODE = "«Code_Postal»"
    CERTIFICATE_ID = "«Numéro_valide»"
    YEAR_N = "«Année n»"
    YEAR_N_1 = "«Année n+1»"
    ID = "«ID»"
    DECHETS_INDUSTRIELS = "«DECHETS_INDUSTRIELS»"
    BIOFUELS_ARTICLE_2 = "«YYYY1»"
    BIOFUELS_ARTICLE_3 = "«YYYY2»"


def application_to_json(application, dechets_industriels="-"):
    year_n = application.period_start.year
    year_n_1 = year_n + 1
    has_dechets_industriels = check_has_dechets_industriels(application)
    biofuels = list(
        application.production.filter(
            year__in=[year_n, year_n_1], feedstock__code__in=["DECHETS_INDUSTRIELS", "AMIDON_RESIDUEL_DECHETS"]
        )
        .values_list("biofuel__name", flat=True)
        .distinct()
    )
    reformatted_biofuels_article_2 = set()
    for biofuel in biofuels:
        if biofuel in name_to_dc_decision_name:
            biofuel = name_to_dc_decision_name.get(biofuel)[1]
        reformatted_biofuels_article_2.add(biofuel)
    reformatted_biofuels_article_2 = list(reformatted_biofuels_article_2)

    reformatted_biofuels_article_3 = set()
    for biofuel in biofuels:
        biofuel_tmp = None
        if biofuel in name_to_dc_decision_name:
            biofuel_tmp = name_to_dc_decision_name.get(biofuel)[2]
        if biofuel_tmp:
            biofuel = biofuel_tmp
        reformatted_biofuels_article_3.add(biofuel)
    reformatted_biofuels_article_3 = list(reformatted_biofuels_article_3)

    return {
        "id": str(application.id),
        "city": application.production_site.city,
        "country": application.production_site.country.name,
        "name": application.production_site.name,
        "certificate_id": application.certificate_id,
        "addresse": application.production_site.address,
        "postal_code": str(application.production_site.postal_code),
        "year_n": str(application.period_start.year),
        "year_n_1": str(application.period_start.year + 1),
        "has_dechets_industriels": has_dechets_industriels,
        "dechets_industriels": dechets_industriels,
        "biofuels_article_2": format_biofuels_to_text(reformatted_biofuels_article_2),
        "biofuels_article_3": format_biofuels_to_text(reformatted_biofuels_article_3),
    }


def format_to_word(data):
    return {
        WordKey.CITY: data.get("city"),
        WordKey.COUNTRY: data.get("country"),
        WordKey.OPERATOR_NAME: data.get("name"),
        WordKey.ADDRESS: data.get("addresse"),
        WordKey.POSTAL_CODE: data.get("postal_code"),
        WordKey.CERTIFICATE_ID: data.get("certificate_id"),
        WordKey.YEAR_N: data.get("year_n"),
        WordKey.YEAR_N_1: data.get("year_n_1"),
        WordKey.ID: data.get("id"),
        WordKey.BIOFUELS_ARTICLE_2: data.get("biofuels_article_2"),
        WordKey.BIOFUELS_ARTICLE_3: data.get("biofuels_article_3"),
        WordKey.DECHETS_INDUSTRIELS: data.get("dechets_industriels").lower(),
    }


def format_biofuels_to_text(biofuels):
    if not biofuels:
        return ""
    if len(biofuels) == 1:
        return f"des {biofuels[0]}".lower()
    elif len(biofuels) == 2:
        return f"des {biofuels[0]} et des {biofuels[1]}".lower()
    else:
        return f"des {', des '.join(biofuels[:-1])} et des {biofuels[-1]}".lower()


def check_has_dechets_industriels(application):
    return application.production.filter(feedstock__code__in=["DECHETS_INDUSTRIELS", "AMIDON_RESIDUEL_DECHETS"]).exists()
