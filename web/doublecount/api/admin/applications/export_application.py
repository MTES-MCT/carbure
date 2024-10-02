import os
from io import BytesIO

from django.conf import settings
from django.db.models import Sum
from django.db.models.query_utils import Q
from django.http import HttpResponse
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from core.common import ErrorResponse
from core.decorators import check_admin_rights
from doublecount.models import (
    DoubleCountingApplication,
)


def replace_and_bold(paragraph, old_text, new_text, font_name="Times New Roman", font_size=12):
    for run in paragraph.runs:
        if old_text in run.text:
            parts = run.text.split(old_text)
            run.text = parts[0]

            new_run = paragraph.add_run(new_text)
            new_run.font.bold = True
            new_run.font.name = font_name
            new_run.font.size = Pt(font_size)
            new_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

            if len(parts) > 1:
                remaining_run = paragraph.add_run(parts[1])
                remaining_run.font.name = font_name
                remaining_run.font.size = Pt(font_size)
                remaining_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_font(paragraph, font_name="Times New Roman", font_size=12, bold=None):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = bold


def set_bold_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True


def set_cell_border(cell, border_color="000000", border_size="4"):
    tc_pr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for border_name in ["top", "bottom", "left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), border_size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color)
        borders.append(border)
    tc_pr.append(borders)


class DoubleCountingApplicationExportError:
    MALFORMED_PARAMS = "MALFORMED_PARAMS"
    APPLICATION_NOT_FOUND = "APPLICATION_NOT_FOUND"
    APPLICATION_NOT_ACCEPTED = "APPLICATION_NOT_ACCEPTED"
    DECHETS_INDUSTRIELS_NOT_FOUND = "DECHETS_INDUSTRIELS_NOT_FOUND"


class WordKey:
    CITY = "«Ville»"
    COUNTRY = "«Pays»"
    OPERATOR_NAME = "«Nom_Opérateur»"
    ADDRESS = "«Adresse»"
    POSTAL_CODE = "«Code_Postal»"
    CERTIFICATE_ID = "«Numéro_valide»"
    YEAR_N = "«Année n»"
    YEAR_N_1 = "«Année n+1»"
    ID = "«ID»"
    DECHETS_INDUSTRIELS = "«DECHETS_INDUSTRIELS»"
    BIOFUELS = "«YYYY»"


def application_to_json(application, dechets_industriels="-"):
    year_n = application.period_start.year
    year_n_1 = year_n + 1
    has_dechets_industriels = check_has_dechets_industriels(application)
    biofuels = list(
        application.production.filter(year__in=[year_n, year_n_1]).values_list("biofuel__name", flat=True).distinct()
    )

    return {
        "id": str(application.id),
        "city": application.production_site.city,
        "country": application.production_site.country.name,
        "name": application.production_site.name,
        "certificate_id": application.certificate_id,
        "addresse": application.production_site.address,
        "postal_code": str(application.production_site.postal_code),
        "year_n": str(application.period_start.year),
        "year_n_1": str(application.period_start.year + 1),
        "has_dechets_industriels": has_dechets_industriels,
        "dechets_industriels": dechets_industriels,
        "biofuels": format_biofuels_to_text(biofuels),
    }


def format_to_word(data):
    return {
        WordKey.CITY: data.get("city"),
        WordKey.COUNTRY: data.get("country"),
        WordKey.OPERATOR_NAME: data.get("name"),
        WordKey.ADDRESS: data.get("addresse"),
        WordKey.POSTAL_CODE: data.get("postal_code"),
        WordKey.CERTIFICATE_ID: data.get("certificate_id"),
        WordKey.YEAR_N: data.get("year_n"),
        WordKey.YEAR_N_1: data.get("year_n_1"),
        WordKey.ID: data.get("id"),
        WordKey.BIOFUELS: data.get("biofuels"),
        WordKey.DECHETS_INDUSTRIELS: data.get("dechets_industriels").lower(),
    }


def format_biofuels_to_text(biofuels):
    if not biofuels:
        return ""
    if len(biofuels) == 1:
        return f"des {biofuels[0]}".lower()
    elif len(biofuels) == 2:
        return f"des {biofuels[0]} et des {biofuels[1]}".lower()
    else:
        return f"des {', des '.join(biofuels[:-1])} et des {biofuels[-1]}".lower()


def check_has_dechets_industriels(application):
    return application.production.filter(feedstock__code="DECHETS_INDUSTRIELS").exists()


@check_admin_rights()
def export_dca(request):
    application_id = request.GET.get("dca_id", False)
    dechets_industriels = request.GET.get("di", "")

    if not application_id:
        return ErrorResponse(400, DoubleCountingApplicationExportError.MALFORMED_PARAMS)

    try:
        application = DoubleCountingApplication.objects.get(id=application_id)
    except Exception:
        return ErrorResponse(400, DoubleCountingApplicationExportError.APPLICATION_NOT_FOUND)

    if application.status != DoubleCountingApplication.ACCEPTED:
        return ErrorResponse(400, DoubleCountingApplicationExportError.APPLICATION_NOT_ACCEPTED)

    has_dechets_industriels = check_has_dechets_industriels(application)
    dechets_industriels = dechets_industriels.strip()
    if "," in dechets_industriels:
        part_1, part_2 = dechets_industriels.rsplit(",", 1)
        dechets_industriels = f"{part_1} et{part_2}"
    if has_dechets_industriels and not dechets_industriels:
        return ErrorResponse(400, DoubleCountingApplicationExportError.DECHETS_INDUSTRIELS_NOT_FOUND)

    if dechets_industriels and not has_dechets_industriels:
        return ErrorResponse(400, DoubleCountingApplicationExportError.MALFORMED_PARAMS)

    if not dechets_industriels:
        dechets_industriels = "-"

    year_n = application.period_start.year
    year_n_1 = year_n + 1

    dcp_data = (
        application.production.filter(year__in=[year_n, year_n_1])
        .values("biofuel__name", "feedstock__name")
        .annotate(
            approved_quota_year_n=Sum("approved_quota", filter=Q(year=year_n)),
            approved_quota_year_n_1=Sum("approved_quota", filter=Q(year=year_n_1)),
        )
    )

    data = application_to_json(application, dechets_industriels)

    to_word = format_to_word(data)

    table_data = [
        {
            **item,
            "approved_quota_year_n_1": (
                "Aucun" if item["approved_quota_year_n_1"] == 0 else item["approved_quota_year_n_1"]
            ),
        }
        for item in dcp_data
    ]

    path = os.path.join(
        settings.BASE_DIR,
        "templates",
        "word",
        "double_count_approval_template.docx",
    )
    doc = Document(path)

    # Tables
    table = doc.tables[1]
    # ... header
    cell = table.rows[0].cells
    cell[2].text = to_word.get(WordKey.YEAR_N)
    cell[3].text = to_word.get(WordKey.YEAR_N_1)

    for cell in table.rows[0].cells:
        set_font(cell.paragraphs[0])

    # ... row
    for item in table_data:
        cell = table.add_row().cells
        cell[0].text = item.get("biofuel__name")
        cell[1].text = item.get("feedstock__name")
        cell[2].text = str(item.get("approved_quota_year_n"))
        cell[3].text = str(item.get("approved_quota_year_n_1"))

        for c in cell:
            set_font(c.paragraphs[0])
            set_cell_border(c)

    # Content
    for index, paragraph in enumerate(doc.paragraphs):
        for k, v in to_word.items():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, v)
                target = "«DECHETS_INDUSTRIELS»"
                if target in paragraph.text:
                    replace_and_bold(paragraph, target, v)

        set_font(paragraph, bold=index in [5, 6, 7, 8])

    # Footer
    section = doc.sections[0]
    footer = section.footer
    for paragraph in footer.paragraphs:
        for k, v in to_word.items():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, v)

        set_font(paragraph)

    # Articles
    found_first_article_3 = False
    ARTICLE_DECHETS_INDUSTRIELS = 3
    current_article_number = 1
    i = 0
    while i < len(doc.paragraphs):
        paragraph = doc.paragraphs[i]

        if paragraph.text.startswith("Article"):
            article_number = int("".join(filter(str.isdigit, paragraph.text)))

            if article_number == ARTICLE_DECHETS_INDUSTRIELS:
                if not has_dechets_industriels and not found_first_article_3:
                    doc.paragraphs[i].clear()
                    doc.paragraphs[i + 1].clear()
                    found_first_article_3 = True
                    i += 1
                else:
                    set_bold_text(paragraph, f"Article {current_article_number}")
                    current_article_number += 1
            else:
                set_bold_text(paragraph, f"Article {current_article_number}")
                current_article_number += 1

        set_font(paragraph)

        i += 1

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f"attachment; filename={application.certificate_id}.docx"

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    response.write(doc_io.read())

    return response
