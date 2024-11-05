# test with : python web/manage.py test doublecount.api.admin.applications.tests_applications.AdminDoubleCountApplicationsTest --keepdb  # noqa: E501
import json
import os
from datetime import date
from io import BytesIO

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from django.urls import reverse
from docx import Document

from certificates.models import DoubleCountingRegistration
from core.models import Entity, MatierePremiere, Pays, UserRights
from core.tests_utils import setup_current_user
from doublecount.api.admin.applications.add import DoubleCountingAddError
from doublecount.api.admin.applications.approve_application import DoubleCountingApplicationApproveError
from doublecount.api.admin.applications.export_application import (
    DoubleCountingApplicationExportError,
    application_to_json,
    check_has_dechets_industriels,
)
from doublecount.factories import (
    DoubleCountingApplicationFactory,
    DoubleCountingProductionFactory,
    DoubleCountingSourcingFactory,
)
from doublecount.factories.agreement import DoubleCountingRegistrationFactory
from doublecount.factories.doc_file import DoubleCountingDocFileFactory
from doublecount.models import DoubleCountingApplication, DoubleCountingDocFile, DoubleCountingProduction
from transactions.models import ProductionSite


class Endpoint:
    change_user_role = reverse("entity-users-change-role")


User = get_user_model()


class AdminDoubleCountApplicationsTest(TestCase):
    fixtures = [
        "json/biofuels.json",
        "json/feedstock.json",
        "json/countries.json",
        "json/depots.json",
        "json/entities.json",
        "json/productionsites.json",
        "json/entities_sites.json",
    ]

    def setUp(self):
        self.admin = Entity.objects.filter(entity_type=Entity.ADMIN)[0]

        self.user = setup_current_user(self, "tester@carbure.local", "Tester", "gogogo", [(self.admin, "RW")], True)

        self.producer, _ = Entity.objects.update_or_create(name="Le Super Producteur 1", entity_type="Producteur")
        UserRights.objects.update_or_create(user=self.user, entity=self.producer, defaults={"role": UserRights.ADMIN})

        self.production_site = ProductionSite.objects.first()
        self.production_site.created_by = self.producer
        self.production_site.address = "1 rue de la Paix"
        france, _ = Pays.objects.update_or_create(code_pays="FR", name="France")
        self.production_site.country = france
        self.production_site.city = "Paris"
        self.production_site.postal_code = "75000"
        self.production_site.dc_reference = ""
        self.production_site.save()
        self.requested_start_year = 2023

    def add_file(self, file_name: str, additional_data=None):
        # upload template
        carbure_home = os.environ["CARBURE_HOME"]
        filepath = f"{carbure_home}/web/fixtures/csv/test_data/{file_name}"
        fh = open(filepath, "rb")
        data = fh.read()
        fh.close()
        f = SimpleUploadedFile("dca.xlsx", data)

        # Add data properties to the post data if provided
        post_data = {
            "entity_id": self.admin.id,
            "producer_id": self.production_site.producer.id,
            "production_site_id": self.production_site.id,
            "file": f,
        }
        if additional_data is not None:
            post_data.update(additional_data)

        response = self.client.post(reverse("admin-double-counting-application-add"), post_data)

        return response

    def create_application(self):
        app = DoubleCountingApplicationFactory.create(
            producer=self.production_site.producer,
            production_site=self.production_site,
            period_start=date(self.requested_start_year, 1, 1),
            period_end=date(self.requested_start_year + 1, 12, 31),
            status=DoubleCountingApplication.PENDING,
        )

        sourcing1 = DoubleCountingSourcingFactory.create(dca=app, year=self.requested_start_year)
        sourcing2 = DoubleCountingSourcingFactory.create(dca=app, year=self.requested_start_year)
        production1 = DoubleCountingProductionFactory.create(
            dca=app, feedstock=sourcing2.feedstock, year=self.requested_start_year + 1, approved_quota=-1
        )
        production2 = DoubleCountingProductionFactory.create(
            dca=app, feedstock=sourcing2.feedstock, year=self.requested_start_year + 1, approved_quota=-1
        )

        return app, sourcing1, production1, sourcing1, production2

    def test_add_application(self):
        # 1 - test add file
        response = self.add_file("dc_agreement_application_valid.xlsx")
        assert response.status_code == 200
        print("**** self.production_site.producer ", self.production_site.__dict__)
        application = DoubleCountingApplication.objects.get(
            producer=self.production_site.producer, period_start__year=self.requested_start_year
        )
        created_at = application.created_at

        # 1.1 - status should be PENDING
        assert application.status == DoubleCountingApplication.PENDING

        # 1.2 - test production requested
        productions = DoubleCountingProduction.objects.filter(dca_id=application.id, year=2023)
        assert productions[0].requested_quota == 20500
        assert productions[1].requested_quota == 10000

        # 1.3  test_dc_number_generation
        assert application.production_site.dc_number == str(1000 + int(application.production_site.id))
        assert application.certificate_id == f"FR_{application.production_site.dc_number}_{self.requested_start_year}"
        assert application.production_site.dc_reference == application.certificate_id

        # 1.2 - check period_start and period_end
        assert application.period_start == date(self.requested_start_year, 1, 1)
        assert application.period_end == date(self.requested_start_year + 1, 12, 31)

        # 2 - test if file uploaded
        DoubleCountingDocFileFactory.create(certificate_id=application.certificate_id, file_name="dca.xlsx")
        docFile = DoubleCountingDocFile.objects.filter(certificate_id=application.certificate_id).first()
        assert docFile.file_name == "dca.xlsx"
        assert docFile.certificate_id == application.certificate_id

        # 3 - test upload twice
        response = self.add_file("dc_agreement_application_valid.xlsx")
        assert response.status_code == 400
        error = response.json()["error"]
        assert error == DoubleCountingAddError.APPLICATION_ALREADY_EXISTS

        # 4 - test should replace
        response = self.add_file("dc_agreement_application_valid.xlsx", {"should_replace": "true"})
        assert response.status_code == 200
        application = DoubleCountingApplication.objects.get(
            producer=self.production_site.producer, period_start__year=self.requested_start_year
        )
        assert application.created_at != created_at

        # 5 - cannot be replaced if already accepted
        application.status = DoubleCountingApplication.ACCEPTED
        application.save()
        response = self.add_file("dc_agreement_application_valid.xlsx", {"should_replace": "true"})
        assert response.status_code == 400
        error = response.json()["error"]
        assert error == DoubleCountingAddError.APPLICATION_ALREADY_RECEIVED

    def test_add_application_to_existing_agreement(self):
        certificate_id = "FR_28_2023"
        # agreement not existing
        response = self.add_file("dc_agreement_application_valid.xlsx", {"certificate_id": certificate_id})
        assert response.status_code == 400
        error = response.json()["error"]
        assert error == DoubleCountingAddError.AGREEMENT_NOT_FOUND

        # agreement existing
        _ = DoubleCountingRegistrationFactory.create(
            certificate_id=certificate_id,
            production_site=self.production_site,
            valid_from=date(self.requested_start_year, 1, 1),
        )

        response = self.add_file("dc_agreement_application_valid.xlsx")
        assert response.status_code == 400
        error = response.json()["error"]
        assert error == DoubleCountingAddError.AGREEMENT_ALREADY_EXISTS

        response = self.add_file("dc_agreement_application_valid.xlsx", {"certificate_id": certificate_id})
        assert response.status_code == 200

        # agreement should be linked to the application
        application = DoubleCountingApplication.objects.get(
            producer=self.production_site.producer, period_start__year=self.requested_start_year
        )
        assert application.certificate_id == certificate_id
        agreement = DoubleCountingRegistration.objects.get(certificate_id=certificate_id)
        assert agreement.application.id == application.id

    def test_production_site_address_mandatory(self):
        self.production_site.address = ""
        self.production_site.save()
        response = self.add_file("dc_agreement_application_valid.xlsx")
        assert response.status_code == 400
        print(response.json())
        error = response.json()["error"]
        assert error == DoubleCountingAddError.PRODUCTION_SITE_ADDRESS_UNDEFINED

    def test_list_applications(self):
        self.create_application()

        response = self.client.get(
            reverse("admin-double-counting-applications"),
            {"entity_id": self.admin.id, "year": self.requested_start_year},
        )

        data = response.json()["data"]
        pending = data["pending"]
        application = pending[0]

        assert application["producer"]["id"] == self.production_site.producer.id

    def test_application_details(self):
        app, sourcing, production, _, _ = self.create_application()

        response = self.client.get(
            reverse("admin-double-counting-application-details"),
            {"entity_id": self.admin.id, "dca_id": app.id},
        )

        application = response.json()["data"]

        assert application["sourcing"][0]["id"] == sourcing.id
        assert application["production"][0]["id"] == production.id

        production_site = application["production_site"]
        assert production_site["id"] == self.production_site.id
        assert production_site["inputs"] == []
        assert production_site["certificates"] == []

    def test_approve_application(self):
        application, sourcing1, production1, sourcing2, production2 = self.create_application()

        # test approve without quotas
        params = {"dca_id": application.id, "entity_id": self.admin.id}
        response = self.client.post(reverse("admin-double-counting-application-approve"), params)
        assert response.status_code == 400
        assert response.json()["error"] == DoubleCountingApplicationApproveError.QUOTAS_NOT_APPROVED
        assert application.status == DoubleCountingApplication.PENDING

        # update quotas
        updated_quotas = [[production1.id, 20500], [production2.id, 10000]]
        response = self.client.post(
            reverse("admin-double-counting-application-update-approved-quotas"),
            {"entity_id": self.admin.id, "approved_quotas": json.dumps(updated_quotas), "dca_id": application.id},
        )
        assert response.status_code == 200

        productions = DoubleCountingProduction.objects.filter(dca_id=application.id)
        assert productions[0].approved_quota == 20500
        assert productions[1].approved_quota == 10000

        application = DoubleCountingApplication.objects.get(
            producer=self.production_site.producer, period_start__year=self.requested_start_year
        )

        # test approve with quotas
        response = self.client.post(reverse("admin-double-counting-application-approve"), params)
        assert response.status_code == 200
        application = DoubleCountingApplication.objects.get(
            producer=self.production_site.producer, period_start__year=self.requested_start_year
        )
        assert application.status == DoubleCountingApplication.ACCEPTED

        # test agreement generation
        agreement = DoubleCountingRegistration.objects.get(certificate_id=application.certificate_id)
        assert agreement.valid_from == date(2023, 1, 1)
        assert agreement.valid_until == date(2024, 12, 31)
        assert agreement.production_site == application.production_site
        assert agreement.certificate_id == application.certificate_id
        assert agreement.application.id == application.id

    def test_export_application(self):
        application, sourcing1, production1, sourcing2, production2 = self.create_application()

        assert application.status != DoubleCountingApplication.ACCEPTED

        # Malformed params
        response = self.client.get(reverse("admin-double-counting-application-export"), {"entity_id": self.admin.id})
        assert response.status_code == 400
        assert response.json()["error"] == DoubleCountingApplicationExportError.MALFORMED_PARAMS

        # Application not found
        response = self.client.get(
            reverse("admin-double-counting-application-export"), {"dca_id": application.id + 200, "entity_id": self.admin.id}
        )
        assert response.status_code == 400
        assert response.json()["error"] == DoubleCountingApplicationExportError.APPLICATION_NOT_FOUND

        # Application not accepted
        response = self.client.get(
            reverse("admin-double-counting-application-export"), {"dca_id": application.id, "entity_id": self.admin.id}
        )
        assert response.status_code == 400
        assert response.json()["error"] == DoubleCountingApplicationExportError.APPLICATION_NOT_ACCEPTED

        application.status = DoubleCountingApplication.ACCEPTED
        application.save()

        assert application.status == DoubleCountingApplication.ACCEPTED

        # Di without di in application
        response = self.client.get(
            reverse("admin-double-counting-application-export"),
            {"dca_id": application.id, "entity_id": self.admin.id, "di": "Graisses brunes, huiles acides"},
        )
        assert response.status_code == 400
        assert response.json()["error"] == DoubleCountingApplicationExportError.MALFORMED_PARAMS

        # Export without di
        assert not check_has_dechets_industriels(application)
        response = self.client.get(
            reverse("admin-double-counting-application-export"), {"dca_id": application.id, "entity_id": self.admin.id}
        )
        assert response.status_code == 200
        assert response["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # ... check doc content
        application_data = application_to_json(application)

        file_stream = BytesIO(response.content)
        doc = Document(file_stream)

        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        full_text = "\n".join(full_text)

        # ... footer
        section = doc.sections[0]
        footer = section.footer

        footer_text = []
        for paragraph in footer.paragraphs:
            footer_text.append(paragraph.text)

        footer_text = "\n".join(footer_text)
        assert application_data["certificate_id"] == footer_text

        # ... content
        del application_data["has_dechets_industriels"]
        del application_data["id"]
        del application_data["dechets_industriels"]

        for _, value in application_data.items():
            assert str(value) in full_text

        assert "Article 5" not in full_text
        assert full_text.count("Article 3") == 1

        # Export with di
        feedstock = MatierePremiere.objects.get(code="DECHETS_INDUSTRIELS")
        assert production1.feedstock.code != "DECHETS_INDUSTRIELS"
        production1.feedstock = feedstock
        production1.save()

        production1.refresh_from_db()

        assert check_has_dechets_industriels(application)

        response = self.client.get(
            reverse("admin-double-counting-application-export"),
            {"dca_id": application.id, "entity_id": self.admin.id, "di": "Graisses brunes, huiles acides"},
        )
        assert response.status_code == 200
        assert response["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        application_data = application_to_json(application)

        file_stream = BytesIO(response.content)
        doc = Document(file_stream)

        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        full_text = "\n".join(full_text)

        # ... content
        del application_data["has_dechets_industriels"]
        del application_data["id"]

        for _, value in application_data.items():
            assert str(value) in full_text

        assert "Article 5" in full_text
        assert full_text.count("Article 3") == 1
